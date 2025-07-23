import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from docx import Document as DocxDocument
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from langchain_core.documents import Document

load_dotenv()
os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))


def extract_text_from_files(files):
    documents = []
    for file in files:
        file_name = file.name
        if file_name.lower().endswith(".pdf"):
            pdf_reader = PdfReader(file)
            for i, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text:
                    doc = Document(
                        page_content=text,
                        metadata={"source": file_name, "page": i + 1}
                    )
                    documents.append(doc)

        elif file_name.lower().endswith((".doc", ".docx")):
            file.seek(0)  # Reset file pointer to the beginning
            docx = DocxDocument(file)
            full_text = "\n".join([para.text for para in docx.paragraphs])
            doc = Document(
                page_content=full_text,
                metadata={"source": file_name, "page": 1}  # word files are treated as one page here
            )
            documents.append(doc)
        else:
            st.warning(f"Unsupported file format: {file_name}")
            
    return documents

 
def get_text_chunks(documents):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_documents(documents)
    return chunks


def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
    vector_store = FAISS.from_documents(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")


def get_conversational_chain():
    prompt_template = """
    Answer the question as detailed as possible from the provided context.
    Make sure to provide all the details. If the answer is not in the provided context,
    just say, "Answer is not available in the provided documents."
    if the anwer has any refrence of image, then provide the image name and page number in the answer.
    **DO NOT include any file names or page numbers in your answer.**
    Only provide the answer based on the text.

    Context:\n {context}\n
    Question: \n{question}\n

    Answer:
    """
    model = ChatGoogleGenerativeAI(model="gemini-2.0-flash", temperature=0.3)

    prompt = PromptTemplate(template = prompt_template, input_variables = ["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

    return chain


def user_input(user_question):
    embeddings = GoogleGenerativeAIEmbeddings(model = "models/embedding-001")
    
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    docs = new_db.similarity_search(user_question)

    chain = get_conversational_chain()
    
    response = chain(
        {"input_documents":docs, "question": user_question}
        , return_only_outputs=True)

    reply_text = response["output_text"]
    return reply_text


def main():
    st.set_page_config("Chat PDF")
    st.header("What information are you looking for?")

    # Initialize chat history in session state
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    # Display chat history
    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    user_question = st.chat_input("Type your question here")

    if user_question:
        # Add user question to chat history
        st.session_state.chat_history.append({"role": "user", "content": user_question})
        with st.chat_message("user"):
            st.markdown(user_question)

        # Get bot response
        bot_response = user_input(user_question)
        
        # Add bot response to chat history
        st.session_state.chat_history.append({"role": "assistant", "content": bot_response})
        with st.chat_message("assistant"):
            st.markdown(bot_response)

    with st.sidebar:
        st.title("Menu:")
        pdf_docs = st.file_uploader("Upload your Files and Click on the Submit & Process Button", accept_multiple_files=True)
        if st.button("Submit & Process"):
            if pdf_docs:
                with st.spinner("Processing..."):
                    documents_with_metadata = extract_text_from_files(pdf_docs)
                    text_chunks = get_text_chunks(documents_with_metadata)
                    get_vector_store(text_chunks)
                    st.success("Done")
                    # Clear chat history when new files are processed
                    st.session_state.chat_history = [] 
            else:
                st.warning("Please upload PDF files first!")


if __name__ == "__main__":
    main()
